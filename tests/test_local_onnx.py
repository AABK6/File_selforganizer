import json
import os
import sys
from pathlib import Path
import subprocess

import pytest
sys.path.append(str(Path(__file__).resolve().parents[1]))
import types
sys.modules.setdefault("onnxruntime_genai", types.SimpleNamespace(Model=object, GeneratorParams=object, Generator=object))
sys.modules.setdefault("transformers", types.SimpleNamespace(AutoTokenizer=types.SimpleNamespace(from_pretrained=lambda *a, **k: types.SimpleNamespace(encode=lambda x: [], decode=lambda x: ""))))

import local_main
import onnx_llm_client


def test_read_txt(tmp_path):
    path = tmp_path / "test.txt"
    path.write_text("hello")
    assert local_main.read_txt(path) == "hello"


def test_read_docx(tmp_path):
    doc_path = tmp_path / "test.docx"
    import docx
    document = docx.Document()
    document.add_paragraph("hello world")
    document.save(doc_path)
    text = local_main.read_docx(doc_path)
    assert "hello world" in text


def test_read_pdf(tmp_path):
    pdf_path = tmp_path / "test.pdf"
    from PyPDF2 import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with open(pdf_path, 'wb') as f:
        writer.write(f)
    # a blank pdf should return an empty string
    assert local_main.read_pdf(pdf_path) == ""


def test_read_doc(monkeypatch):
    def fake_run(args, capture_output, text):
        class R:
            returncode = 0
            stdout = "dummy"
        return R()
    monkeypatch.setattr(subprocess, "run", fake_run)
    assert local_main.read_doc(Path("a.doc")) == "dummy"


def test_read_doc_no_antiword(monkeypatch):
    def fake_run(args, capture_output, text):
        raise FileNotFoundError
    monkeypatch.setattr(subprocess, "run", fake_run)
    assert local_main.read_doc(Path("a.doc")) == ""


def test_read_file_dispatch(tmp_path, monkeypatch):
    txt = tmp_path / "a.txt"
    txt.write_text("hi")
    assert local_main.read_file(txt) == "hi"

    docx_path = tmp_path / "b.docx"
    import docx as docx_module
    d = docx_module.Document()
    d.add_paragraph("hi")
    d.save(docx_path)
    assert "hi" in local_main.read_file(docx_path)

    pdf = tmp_path / "c.pdf"
    from PyPDF2 import PdfWriter
    w = PdfWriter()
    w.add_blank_page(width=72, height=72)
    with open(pdf, 'wb') as f:
        w.write(f)
    assert local_main.read_file(pdf) == ""


def test_onnx_client_analyze(monkeypatch, tmp_path):
    model_dir = tmp_path / "model"
    model_dir.mkdir()

    def fake_init(self, model_dir):
        self.model_dir = model_dir
    monkeypatch.setattr(onnx_llm_client.ONNXLLMClient, "__init__", fake_init)

    client = onnx_llm_client.ONNXLLMClient(str(model_dir))
    monkeypatch.setattr(client, "_generate", lambda prompt, max_tokens=1024: '[{"k":1}]')
    assert client.analyze_text("text") == [{"k": 1}]


def test_onnx_client_propose(monkeypatch, tmp_path):
    model_dir = tmp_path / "model"
    model_dir.mkdir()

    def fake_init(self, model_dir):
        self.model_dir = model_dir
    monkeypatch.setattr(onnx_llm_client.ONNXLLMClient, "__init__", fake_init)

    client = onnx_llm_client.ONNXLLMClient(str(model_dir))
    monkeypatch.setattr(client, "_generate", lambda prompt, max_tokens=1024: '{"folder": []}')
    assert client.propose_structure({}) == {"folder": []}


def test_main_integration(tmp_path, monkeypatch):
    input_dir = tmp_path / "files"
    input_dir.mkdir()
    (input_dir / "f.txt").write_text("hello")
    model_dir = tmp_path / "model"
    model_dir.mkdir()

    class FakeClient:
        def __init__(self, md):
            pass
        def analyze_text(self, text):
            return {"summary": "ok"}
        def propose_structure(self, analysis):
            return {"folder": list(analysis.keys())}

    monkeypatch.setattr(local_main, "ONNXLLMClient", FakeClient)

    cwd = os.getcwd()
    os.chdir(tmp_path)
    try:
        sys.argv = ["local_main.py", str(input_dir), "--model_dir", str(model_dir)]
        local_main.main()
    finally:
        os.chdir(cwd)
    report = tmp_path / "report.json"
    assert report.exists()
    data = json.loads(report.read_text())
    assert "analysis" in data and "structure" in data

